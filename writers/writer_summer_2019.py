from docx import Document
from os import listdir
from docx.shared import Pt

PATTERN_FILENAME = 'pattern/pattern_summer_2019.docx'

def replace(source, course_info, prices_info):
		source = source.replace('COURSE_NAME', str(course_info.course_name))
		source = source.replace('AUTHORNAME', str(course_info.author_name))
		source = source.replace('AUTHOR_NAME', str(course_info.author_name))
		source = source.replace('CONTRACT_PRICE_RUB_IN_WORDS', str(course_info.contract_price_rub_in_words))
		source = source.replace('CONTRACT_PRICE_CENT_IN_WORDS', str(course_info.contract_price_cent_in_words))
		source = source.replace('CONTRACT_PRICE_RUB', str(course_info.contract_price_rub))
		source = source.replace('CONTRACT_PRICE_CENT', str(course_info.contract_price_cent))
		source = source.replace('AUTHOR_REWARD_RUB_IN_WORDS', str(course_info.author_reward_rub_in_words))
		source = source.replace('AUTHOR_REWARD_CENT_IN_WORDS', str(course_info.author_reward_cent_in_words))
		source = source.replace('AUTHOR_REWARD_RUB', str(course_info.author_reward_rub))
		source = source.replace('AUTHOR_REWARD_CENT', str(course_info.author_reward_cent))
		source = source.replace('INSURANCE_RUB_IN_WORDS', str(course_info.insurance_rub_in_words))
		source = source.replace('INSURANCE_CENT_IN_WORDS', str(course_info.insurance_cent_in_words))
		source = source.replace('INSURANCE_RUB', str(course_info.insurance_rub))
		source = source.replace('INSURANCE_CENT', str(course_info.insurance_cent))
		source = source.replace('ADDRESS_BY_PASPORT', str(course_info.address_by_pasport))
		source = source.replace('TELEPHONE_NUMBER', str(course_info.telephone_number))
		source = source.replace('AUTHOR_SECOND_NAME', str(course_info.author_second_name))
		source = source.replace('AUTHOR_FIRST_NAME', str(course_info.author_first_name))
		source = source.replace('AUTHOR_MIDDLE_NAME', str(course_info.author_middle_name))
		source = source.replace('AUTHOR_BIRTH_DATE', str(course_info.birthdate))
		source = source.replace('PASPORT_SERIES_NUMBER', str(course_info.passport_seria_number))
		source = source.replace('PASPORT_CREATED_LOCATION', str(course_info.passport_created_location))
		source = source.replace('PASPORT_CREATED_DATE', str(course_info.passport_created_date))
		source = source.replace('BANK_NAME', str(course_info.bank_name))
		source = source.replace('JOB_NAME', str(course_info.job_name))
		source = source.replace('ITN_10', str(course_info.itn_10))
		source = source.replace('ITN_11', str(course_info.itn_11))
		source = source.replace('ITN_12', str(course_info.itn_12))
		source = source.replace('ITN_1', str(course_info.itn_1))
		source = source.replace('ITN_2', str(course_info.itn_2))
		source = source.replace('ITN_3', str(course_info.itn_3))
		source = source.replace('ITN_4', str(course_info.itn_4))
		source = source.replace('ITN_5', str(course_info.itn_5))
		source = source.replace('ITN_6', str(course_info.itn_6))
		source = source.replace('ITN_7', str(course_info.itn_7))
		source = source.replace('ITN_8', str(course_info.itn_8))
		source = source.replace('ITN_9', str(course_info.itn_9))
		source = source.replace('ITN', str(course_info.ITN))
		source = source.replace('INILA_10', str(course_info.inila_10))
		source = source.replace('INILA_11', str(course_info.inila_11))
		source = source.replace('INILA_1', str(course_info.inila_1))
		source = source.replace('INILA_2', str(course_info.inila_2))
		source = source.replace('INILA_3', str(course_info.inila_3))
		source = source.replace('INILA_4', str(course_info.inila_4))
		source = source.replace('INILA_5', str(course_info.inila_5))
		source = source.replace('INILA_6', str(course_info.inila_6))
		source = source.replace('INILA_7', str(course_info.inila_7))
		source = source.replace('INILA_8', str(course_info.inila_8))
		source = source.replace('INILA_9', str(course_info.inila_9))
		source = source.replace('INILA', str(course_info.INILA))
		source = source.replace('TAX_START_DATE', str(course_info.tax_start_date))
		source = source.replace('POSITION', str(course_info.position))
		source = source.replace('BANK_REQUISITES', str(course_info.bank_requisites))
		source = source.replace('DEGREE', str(course_info.degree))

		source = source.replace('LECTURES_HOURS', str(prices_info.lectures_hours))
		source = source.replace('LECTURES_PRICE', str(prices_info.lectures_price))
		source = source.replace('LECTURES_COST', str(prices_info.lectures_cost))
		source = source.replace('STRUCTURE_PRICE', str(prices_info.structure_price))
		source = source.replace('STRUCTURE_HOURS', str(prices_info.structure_hours))
		source = source.replace('STRUCTURE_COST', str(prices_info.structure_cost))
		source = source.replace('TESTS_PRICE', str(prices_info.tests_price))
		source = source.replace('TESTS_HOURS', str(prices_info.tests_hours))
		source = source.replace('TESTS_COST', str(prices_info.tests_cost))
		source = source.replace('GUIDANCE_PRICE', str(prices_info.guidance_price))
		source = source.replace('GUIDANCE_HOURS', str(prices_info.guidance_hours))
		source = source.replace('GUIDANCE_COST', str(prices_info.guidance_cost))

		return source
		
def write_data_to_file(course_info, prices_info):

		document = Document(PATTERN_FILENAME)
		for p in document.paragraphs:
			p.text = replace(p.text, course_info, prices_info)			
			p.style.font.name = 'Times New Roman'
			p.style.font.size = Pt(11)

		for t in document.tables:
			for row_number in range(len(t.column_cells(0))):
				for cell in t.row_cells(row_number):
					cell.text = replace(cell.text, course_info, prices_info)

		file_name = course_info.author_second_name + ' - ' \
			+ course_info.course_name + '.docx'
		if len(course_info.passport_seria_number) > 5:
			folder_name = 'results/'
		else:
			folder_name = 'no-passport/'
		while file_name in listdir(folder_name):
			file_name = file_name.replace('.docx', '1.docx')
		document.save(folder_name + file_name)